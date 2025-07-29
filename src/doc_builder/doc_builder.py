from docx import Document, types
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from event_calendar.event_calendar import EventCalendar


class DocBuilder:
    """
    The DocBuilder class describes a Docx based document builder that will be used to format the calendar page for printing.

      Constructor:
        The constructor will create an instance of the Document class from the python-docx library, then initialize some settings that are imported from the settings module. Since the calendar is a singular file and source of information that won't be changed, this information can all be satic. Any style changes should be made in the settings module.
    """

    def __init__(self, font_style: str, font_size: int, margins: dict, calendar: EventCalendar = None):
        self.__doc = Document()
        self.__table = None
        self.calendar = calendar
        self.margins = margins
        self.set_page_orientation()
        self.set_doc_styles(font_style, font_size)
        self.set_page_header()

    @property
    def margins(self) -> dict:
        """The margin sizes of the document in inches."""
        return self.__margins

    @margins.setter
    def margins(self, margins: dict) -> dict:
        """
        Sets the margins for the document.

          Parameters:
              margins {dict} —— margin values in inches with keys "top", "bottom", "left", and "right"
          Raises:
              KeyError if the margins dictionary does not contain the required keys.
              AttributeError if the object does not have the required margin attributes.
          Returns:
            {dict} containing the margin values in inches if successful.
        """

        if not isinstance(margins, dict):
            raise TypeError("Margins must be a dictionary.")

        for key in ["top", "bottom", "left", "right"]:
            if key not in margins:
                raise KeyError(f"Margins dictionary must contain '{key}' key.")
            if margins[key] < 0:
                raise ValueError(f"Margin '{key}' cannot be negative.")

        try:
            sections = self.__doc.sections
            sections.left_margin = Inches(margins["left"])
            sections.right_margin = Inches(margins["right"])
            sections.top_margin = Inches(margins["top"])
            sections.bottom_margin = Inches(margins["bottom"])
            self.__margins = margins

        except AttributeError as e:
            print(
                f"The object does not have the required margin attributes: {e}")
            raise

    @property
    def font_style(self) -> str:
        """The font style of the document."""
        return self.__font_style

    @font_style.setter
    def font_style(self, style: str) -> str:
        """
        Sets the default font style for the document.

          Parameters:
              font_style {str} ——  the font style to set
          Raises:
            AttributeError if the object does not have a font style attribute.
          Returns:
              {str} —— the font style as a  if successful.
        """

        if not isinstance(style, str):
            raise TypeError("Font style must be a string.")
        try:
            self.__font_style = style
        except AttributeError as e:
            print(
                f"The object does not have the required style attribute: {e}")
            raise

    @property
    def font_size(self) -> int:
        """The font size of the document."""
        return self.__font_size

    @font_size.setter
    def font_size(self, size: int) -> int:
        """
        Sets the default font size for the document.

          Parameters:
              font_size {int} —— the font size to set
          Raises:
            TypeError if the font size is not an integer.
            AttributeError if the object does not have a font size attribute.
          Returns:
              {int} —— The font size if successful.
        """

        if not isinstance(size, int):
            raise TypeError("Font size must be an integer.")
        try:
            self.__font_size = size
        except AttributeError as e:
            print(
                f"The object does not have the required size attribute: {e}")
            raise

    @property
    def styles(self):
        return self.__doc.styles

    def set_doc_styles(self, font_style: str, font_size: int) -> None:
        """
        Sets the font style and size for the whole document

          Parameters:
            font_style {str} —— the font style to set
            font_size {int} —— the font size to set

          Returns:
              None
        """
        style = self.__doc.styles["Normal"]
        font = style.font
        font.name = font_style
        font.size = font_size

    def set_page_orientation(self) -> None:
        """
        Sets the document up as a default 8.5 x 11 landscape page.
        Returns: None
        """

        main_section = self.__doc.sections[-1]
        main_section.orientation = WD_ORIENT.LANDSCAPE
        main_section.page_width = Inches(11.0)
        main_section.page_height = Inches(8.5)

    def set_page_header(self):
        top_section = self.__doc.sections[0]
        header = top_section.header
        text = header.paragraphs[0]
        text.alignment = WD_TABLE_ALIGNMENT.CENTER
        text.style.font.size = Pt(20)
        text.bold = True
        text.text = "Meeting Room Schedule"

    def init_table(self):
        self.__table = self.create_table(
            rows=3, cols=6, cell_width_inches=1.65)
        self.__table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.__table.autofit = False
        self.__table.style = 'Table Grid'
        self.set_table_weekdays()
        self.set_table_dates()
        self.populate_table_with_events()

    def create_table(self, rows, cols, cell_width_inches, parent=None) -> types.ProvidesXmlPart:
        """
        Creates a table with specified dimensions and appends it to parent. If no parent is supplied, appends directly to the main document.

        Parameters:
            rows ——
            cols ——
            cell_width_inches ——
            parent ——
        Returns:
            table —— The created table
        """
        if not parent:
            parent = self.__doc
        table = parent.add_table(rows=rows, cols=cols)

        # set widths of the table columns
        cols = table.columns
        for i in range(0, len(cols)):
            for cell in range(0, len(cols[i].cells)):
                cols[i].cells[cell].width = Inches(cell_width_inches)

        return table

    def set_table_weekdays(self):
        # set headers of the table
        day_headers = self.__table.rows[0].cells
        week = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat"]
        for day in range(0, len(week)):
            paragraph = day_headers[day].paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            paragraph.alignment - WD_ALIGN_VERTICAL.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            text = paragraph.add_run(week[day])
            text.bold = True
            text.font.color.rgb = RGBColor(255, 255, 255)
            text.font.size = Pt(12)

            shading = parse_xml(
                r'<w:shd {} w:fill="5b9bd7"/>'.format(nsdecls('w')))
            styles = day_headers[day]._element.get_or_add_tcPr()
            styles.append(shading)

    def set_table_dates(self):
        # set dates of the table
        dates = self.calendar.get_next_weeks_dates()
        date_headers = self.__table.rows[1].cells
        for date in range(0, len(dates)):
            paragraph = date_headers[date].paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.RIGHT
            text = paragraph.add_run(dates[date][-2::])

    def populate_table_with_events(self):
        # set table events
        event_containers = self.__table.rows[2].cells
        for index, (__, events) in enumerate(self.calendar.events.items()):
            container = event_containers[index]

            cell_idx = 0
            for event in events:
                table = self.create_table(1, 1, 1.4, container)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # style cell
                cell = table.cell(0, 0)
                styles = cell._element.get_or_add_tcPr()
                borders = OxmlElement('w:tcBorders')
                styles.append(borders)
                self.append_border_styles(borders)

                # add data to cell
                event_text = cell.paragraphs[0]
                event_text.paragraph_format.space_before = Pt(0)
                title = event_text.add_run(event.title + "\n")
                title.bold = True
                date = event_text.add_run(event.full_event_string())
                cell_idx += 1

    def append_border_styles(self, border, sides=['top', 'left', 'bottom', 'right']):
        for border_type in sides:
            border_elm = OxmlElement(f'w:{border_type}')
            border_elm.set(qn('w:val'), 'single')
            border_elm.set(qn('w:sz'), '10')
            border_elm.set(qn('w:space'), '0')
            border_elm.set(qn('w:color'), '000000')
            border.append(border_elm)

    def init_psa(self):
        quiet_room = self.create_table(1, 1, 9.9)
        quiet_room.alignment = WD_TABLE_ALIGNMENT.CENTER
        quiet_room.autofit = False
        quiet_room.rows[0].height = Inches(2)

        container = quiet_room.rows[0].cells[0]
        paragraph = container.paragraphs[0]
        paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

        styles = self.styles
        psa_style = styles.add_style('psa', WD_STYLE_TYPE.CHARACTER)
        psa_font = psa_style.font
        psa_font.color.rgb = RGBColor(255, 255, 255)
        psa_font.size = Pt(30)
        psa_text = "The meeting room is available for public use as a quiet space when not reserved."
        paragraph.add_run(psa_text, style='psa').bold = True

        # style cell
        cell = quiet_room.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        styles = cell._element.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        styles.append(borders)
        self.append_border_styles(borders, ['left', 'bottom', 'right'])

        shading = parse_xml(
            r'<w:shd {} w:fill="5b9bd7"/>'.format(nsdecls('w')))
        styles = container._element.get_or_add_tcPr()
        styles.append(shading)

    def save_document(self, file_path):
        self.__doc.save(file_path)
